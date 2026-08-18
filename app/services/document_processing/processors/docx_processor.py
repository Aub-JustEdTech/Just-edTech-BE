"""DOCX document processor"""

import logging
from typing import Any

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    Document = None
    qn = None
    Table = None
    Paragraph = None

from app.services.document_processing.base import DocumentProcessor
from app.services.document_processing.markdown_table import (
    cell_value,
    make_separator,
    row_to_markdown,
)

logger = logging.getLogger(__name__)

# python-docx style names for heading-like paragraphs, mapped to a nesting
# level. Custom/localized styles (e.g. French "Titre 1") aren't recognized —
# such documents degrade gracefully to the no-headings case in
# `extract_sections` rather than erroring.
_TITLE_STYLES = {"title": 1, "subtitle": 2}


def _heading_level(style_name: str) -> int | None:
    """Return the nesting level for a heading/title style, else None."""
    name = (style_name or "").strip().lower()
    if name in _TITLE_STYLES:
        return _TITLE_STYLES[name]
    if name.startswith("heading "):
        suffix = name[len("heading ") :].strip()
        if suffix.isdigit():
            return int(suffix)
    return None


def _iter_block_items(document):
    """
    Yield each top-level `Paragraph`/`Table` in document order.

    python-docx's `.paragraphs` and `.tables` are separate, order-losing
    collections — this walks the underlying XML body directly so paragraphs
    and tables interleave the way they actually appear in the document.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_to_markdown(table) -> str:
    """Render a DOCX table as Markdown, treating its first row as a header."""
    rows = [[cell_value(cell.text) for cell in row.cells] for row in table.rows]
    rows = [r for r in rows if any(v for v in r)]
    if not rows:
        return ""

    header, *data_rows = rows
    col_count = len(header)
    lines = [row_to_markdown(header), make_separator(col_count)]
    for row in data_rows:
        padded = (row + [""] * col_count)[:col_count]
        lines.append(row_to_markdown(padded))
    return "\n".join(lines)


class DocxProcessor(DocumentProcessor):
    """Process DOCX (Microsoft Word) documents"""

    supported_extensions = [".docx"]
    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX processing. Install it with: pip install python-docx"
            )

        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise

    def extract_metadata(self, file_path: str) -> dict[str, Any]:
        """Extract DOCX metadata"""
        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX processing. Install it with: pip install python-docx"
            )

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            # Count paragraphs and tables
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)

            # Get text for word count
            text = "\n".join([p.text for p in doc.paragraphs])
            word_count = len(text.split())

            metadata = {
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "word_count": word_count,
                "character_count": len(text),
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
            return metadata
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {"paragraph_count": 0, "table_count": 0}

    def validate(self, file_path: str) -> bool:
        """Validate DOCX file"""
        if Document is None:
            return False

        try:
            Document(file_path)
            return True
        except Exception:
            return False

    def extract_sections(self, file_path: str) -> list[dict[str, Any]]:
        """
        Walk the document in order, grouping body text and tables under the
        heading path they fall under.

        Returns a list of dicts, in document order:
            {
                "heading_path": "Board Policy > Section 3.2",  # "" if none
                "text": <section body text, or a rendered Markdown table>,
                "is_table": <True for a table section, else False>,
            }

        A document with no headings at all produces a single section with
        `heading_path=""` — callers (see `heading_chunker.py`) sub-chunk an
        oversized section the same way regardless of whether it came from a
        real heading or this no-headings fallback.
        """
        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX processing. Install it with: pip install python-docx"
            )

        doc = Document(file_path)
        sections: list[dict[str, Any]] = []
        stack: list[tuple[int, str]] = []  # (level, heading_text)
        buffer: list[str] = []

        def heading_path() -> str:
            return " > ".join(h for _, h in stack)

        def flush() -> None:
            text = "\n\n".join(buffer).strip()
            buffer.clear()
            if text:
                sections.append(
                    {"heading_path": heading_path(), "text": text, "is_table": False}
                )

        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                level = _heading_level(block.style.name if block.style else "")
                if level is not None:
                    # A new heading always starts a new section, even if the
                    # current buffer is empty (heading immediately followed
                    # by another heading) — `flush()` is a no-op on empty
                    # buffers, so the parent heading text simply folds into
                    # the child's heading_path with no special case needed.
                    flush()
                    while stack and stack[-1][0] >= level:
                        stack.pop()
                    stack.append((level, block.text.strip()))
                elif block.text.strip():
                    buffer.append(block.text)
            elif isinstance(block, Table):
                flush()
                table_md = _table_to_markdown(block)
                if table_md:
                    sections.append(
                        {
                            "heading_path": heading_path(),
                            "text": table_md,
                            "is_table": True,
                        }
                    )

        flush()
        return sections
