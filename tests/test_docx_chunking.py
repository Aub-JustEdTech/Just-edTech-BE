"""Tests for DocxProcessor.extract_sections and the shared heading-chunk
packing helper. The regression under test: python-docx's `.paragraphs` and
`.tables` are separate, order-losing collections, which is why the old
`extract_text` flattened tables into the paragraph stream out of document
order. `extract_sections` walks the document body directly to fix that.
"""

from __future__ import annotations

from docx import Document

from app.services.document_processing.heading_chunker import chunk_sections_by_heading
from app.services.document_processing.processors.docx_processor import DocxProcessor


def _write_docx(tmp_path, build) -> str:
    doc = Document()
    build(doc)
    path = tmp_path / "doc.docx"
    doc.save(path)
    return str(path)


def test_nested_headings_produce_correct_heading_path(tmp_path):
    def build(doc):
        doc.add_heading("Board Policy", level=1)
        doc.add_paragraph("Intro text under the top heading.")
        doc.add_heading("Section 3.2", level=2)
        doc.add_paragraph("Body text under the nested heading.")

    path = _write_docx(tmp_path, build)
    sections = DocxProcessor().extract_sections(path)

    paths = [s["heading_path"] for s in sections]
    assert "Board Policy" in paths
    assert "Board Policy > Section 3.2" in paths


def test_table_between_paragraphs_kept_in_document_order(tmp_path):
    def build(doc):
        doc.add_paragraph("Paragraph before the table.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Vote"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "Yes"
        doc.add_paragraph("Paragraph after the table.")

    path = _write_docx(tmp_path, build)
    sections = DocxProcessor().extract_sections(path)

    kinds = [(s["is_table"], s["text"]) for s in sections]
    assert len(kinds) == 3
    assert kinds[0][0] is False and "before the table" in kinds[0][1]
    assert kinds[1][0] is True and "Alice" in kinds[1][1]
    assert kinds[2][0] is False and "after the table" in kinds[2][1]


def test_docx_with_no_headings_produces_one_section(tmp_path):
    def build(doc):
        doc.add_paragraph("Just some plain text.")
        doc.add_paragraph("And some more plain text.")

    path = _write_docx(tmp_path, build)
    sections = DocxProcessor().extract_sections(path)

    assert len(sections) == 1
    assert sections[0]["heading_path"] == ""
    assert "plain text" in sections[0]["text"]


def test_heading_immediately_followed_by_subheading_has_no_empty_section(tmp_path):
    def build(doc):
        doc.add_heading("Section 3", level=1)
        doc.add_heading("3.1 Overview", level=2)
        doc.add_paragraph("Only the subsection has body text.")

    path = _write_docx(tmp_path, build)
    sections = DocxProcessor().extract_sections(path)

    assert len(sections) == 1
    assert sections[0]["heading_path"] == "Section 3 > 3.1 Overview"


def test_empty_docx_produces_no_sections(tmp_path):
    path = _write_docx(tmp_path, lambda doc: None)
    sections = DocxProcessor().extract_sections(path)
    assert sections == []


# ---------------------------------------------------------------------------
# chunk_sections_by_heading — shared packing helper
# ---------------------------------------------------------------------------


def test_oversized_section_is_subsplit_preserving_heading_path():
    long_text = "This is one sentence. " * 200
    sections = [{"heading_path": "Long Section", "text": long_text, "is_table": False}]

    packed = chunk_sections_by_heading(sections, chunk_size=200, chunk_overlap=20)

    assert len(packed) > 1
    for chunk in packed:
        assert chunk["heading_path"] == "Long Section"
        assert len(chunk["text"]) <= 200


def test_sections_are_never_merged_across_different_headings():
    sections = [
        {"heading_path": "A", "text": "Short A.", "is_table": False},
        {"heading_path": "B", "text": "Short B.", "is_table": False},
    ]
    packed = chunk_sections_by_heading(sections, chunk_size=1000, chunk_overlap=100)

    assert len(packed) == 2
    assert packed[0]["heading_path"] == "A"
    assert packed[1]["heading_path"] == "B"


def test_empty_section_is_skipped():
    sections = [
        {"heading_path": "A", "text": "   ", "is_table": False},
        {"heading_path": "B", "text": "Real content.", "is_table": False},
    ]
    packed = chunk_sections_by_heading(sections, chunk_size=1000, chunk_overlap=100)

    assert len(packed) == 1
    assert packed[0]["heading_path"] == "B"
